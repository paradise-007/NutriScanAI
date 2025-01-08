from docx import Document
from docx.shared import Inches, RGBColor
from Database import retrieve_data, retrive_count
from Graph import pie_graph, line_graph, stacked_bar_graph
from Home import url, db, user_collection
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from io import BytesIO

# Set font color to blue
gender_options = retrieve_data(url, db, user_collection, field='gender')
state_options = retrieve_data(url, db, user_collection, field='state')

gender_count = np.empty(len(gender_options))
state_count = np.empty(len(state_options))
total_user = retrive_count(url, db, user_collection, {})

for i in range(len(gender_options)):
    obj = {'gender': gender_options[i]}
    gender_count[i] = int(retrive_count(url, db, user_collection, obj))

for i in range(len(state_options)):
    obj = {'state': state_options[i]}
    state_count[i] = int(retrive_count(url, db, user_collection, obj))

# Create DataFrames for gender and state counts
gender_df = pd.DataFrame({'Gender': gender_options, 'Count': gender_count})
state_df = pd.DataFrame({'State': state_options, 'Count': state_count})

def generate_report(username, past_days, figs):
    document = Document()

    # Title
    title = document.add_heading('NutriScanAI Report', level=0)
    title.alignment = 1  # Center the title
    title_run = title.runs[0]
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 255)  # Set font color to blue

    # Greeting message
    greeting = document.add_paragraph(f'Hello {username},\n\n')
    greeting.alignment = 1  # Center the greeting

    # Total user count
    document.add_paragraph(f'Total Users: {total_user}', style='Normal')

    # Gender Count Table
    document.add_heading('Gender Distribution', level=1)
    gender_table = document.add_table(rows=1, cols=2)
    hdr_cells = gender_table.rows[0].cells
    hdr_cells[0].text = 'Gender'
    hdr_cells[1].text = 'Count'

    for index, row in gender_df.iterrows():
        row_cells = gender_table.add_row().cells
        row_cells[0].text = row['Gender']
        row_cells[1].text = str(row['Count'])

    # State Count Table
    document.add_heading('State Wise User Count', level=1)
    state_table = document.add_table(rows=1, cols=4)
    hdr_cells = state_table.rows[0].cells
    hdr_cells[0].text = 'State'
    hdr_cells[1].text = 'Male Users'
    hdr_cells[2].text = 'Female Users'
    hdr_cells[3].text = 'Total Users'

    for index, row in state_df.iterrows():
        state = row['State']
        male_count = int(retrive_count(url, db, user_collection, {'state': state, 'gender': 'Male'}))
        female_count = int(retrive_count(url, db, user_collection, {'state': state, 'gender': 'Female'}))
        total_count = male_count + female_count

        row_cells = state_table.add_row().cells
        row_cells[0].text = state
        row_cells[1].text = str(male_count)
        row_cells[2].text = str(female_count)
        row_cells[3].text = str(total_count)

    # Graphs generation
    fig1 = line_graph(x_axis='date', y_axis='gender', days=past_days)
    fig2 = pie_graph(variable='gender', top=2, color=['pink'], days=past_days)
    fig3 = pie_graph(variable='state', top=5, days=past_days)
    fig4 = pie_graph(count_types=True, days=past_days)
    fig5 = line_graph(days=past_days, total_users=True)
    fig6 = stacked_bar_graph(x_axis='state', y_axis='gender', days=past_days, top=5)

    # Standardized figure size
    figure_width = 4
    figure_height = 2

    # Helper function to add graphs
    def add_graph_to_document(fig):
        img_stream = BytesIO()
        fig.savefig(img_stream, format='png', bbox_inches='tight')
        img_stream.seek(0)

        # Create a table for the image with borders
        table = document.add_table(rows=1, cols=1)
        cell = table .cell(0, 0)
        run = cell.add_paragraph().add_run()
        run.add_picture(img_stream, width=Inches(figure_width))
        cell.width = Inches(figure_width + 1)  # Adjust cell width for borders

        # Center the table
        table.alignment = 1  # Center the table

    # Add all graphs to the document
    for fig in figs:
        add_graph_to_document(fig)

    # Save document
    document.save('NutriScanAI_Report.docx')